# -*- coding: utf-8 -*-
"""Week 1 Internship Report generator - built on Spurvance Labs template (keeps logo)."""

import os
import glob
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ---- STUDENT INFO ----
STUDENT_NAME = "Gul Shair"
DEPARTMENT = "AI Engineer"
SUPERVISOR = "Professor Sahib"
SUBMISSION_DATE = "26 June 2026"
TEMPLATE = "report_template.docx"
SCREENSHOTS_DIR = "screenshots"
OUT = "GulShair_AIEngineer_Week1_Report.docx"
# ----------------------

doc = Document(TEMPLATE)  # keeps header logo + footer
ACCENT = RGBColor(0x1F, 0x4E, 0x79)


def fill_details_table():
    """Fill the existing Student Details table values (column 2)."""
    values = {
        "name": STUDENT_NAME,
        "department": DEPARTMENT,
        "supervisor": SUPERVISOR,
        "submission": SUBMISSION_DATE,
    }
    for t in doc.tables:
        for row in t.rows:
            label = row.cells[0].text.strip().lower()
            if "name" == label or label.startswith("name"):
                if "supervisor" not in label:
                    row.cells[1].text = values["name"]
            if "department" in label:
                row.cells[1].text = values["department"]
            if "supervisor" in label:
                row.cells[1].text = values["supervisor"]
            if "submission" in label:
                row.cells[1].text = values["submission"]


def heading(text, size=14):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = ACCENT
    return p


def body(text):
    doc.add_paragraph(text)


def bullet(text):
    from docx.shared import Inches as _In
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = _In(0.3)
    p.paragraph_format.space_after = Pt(2)
    p.add_run("\u2022  " + text)


def make_table(rows):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    for sname in ("Table Grid", "Light Grid Accent 1", "Light List Accent 1"):
        try:
            t.style = sname
            break
        except Exception:
            continue
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            t.rows[i].cells[j].text = val
            if i == 0 and t.rows[i].cells[j].paragraphs[0].runs:
                t.rows[i].cells[j].paragraphs[0].runs[0].bold = True
    doc.add_paragraph()
    return t


# ===== Fill the existing Student Details table (under the logo) =====
fill_details_table()

# A small spacer before the report content
doc.add_paragraph()

# ---------------- 1. Week Objective ----------------
heading("1. Week Objective")
body(
    "The objective of Week 1 was to understand the fundamentals of building a client-server "
    "application for audio handling and to begin research into Text-to-Speech (TTS) and voice "
    "cloning models. The practical goal was to set up the development environment and build a "
    "working Audio File Transfer application, using a FastAPI backend to receive and send audio "
    "files and a Streamlit frontend as the user interface. Alongside the build, the week was used "
    "to research suitable open-source TTS and voice cloning models that will form the core of the "
    "project in the coming weeks."
)

# ---------------- 2. Tasks Completed ----------------
heading("2. Tasks Completed")

heading("2.1  Environment and Project Setup", 12)
bullet("Installed Python and created a virtual environment (.venv) to keep the project dependencies isolated.")
bullet("Created the project structure with separate files for the backend (server.py) and frontend (app.py).")
bullet("Created requirements.txt listing all required libraries: fastapi, uvicorn[standard], python-multipart, streamlit, and requests.")
bullet("Installed all dependencies inside the virtual environment using pip install -r requirements.txt.")

heading("2.2  FastAPI Backend (server.py)", 12)
bullet("Built a FastAPI server titled \"Audio Transfer Server\" to handle audio file transfer.")
bullet("Created a health-check endpoint (GET /) that returns the server running status.")
bullet("Implemented an upload endpoint (POST /receive) that accepts an uploaded audio file and saves it into the uploads/ folder using shutil.")
bullet("Implemented a download endpoint (GET /send) that returns the most recently uploaded file using FastAPI's FileResponse, with a clean 404 response when no file is available.")
bullet("Used os.makedirs to auto-create the uploads/ storage folder on startup so files are stored safely.")

heading("2.3  Streamlit Frontend (app.py)", 12)
bullet("Built a clean Streamlit UI titled \"Audio File Transfer\" with two sections: Upload and Download.")
bullet("Section 1 lets the user choose an audio file (.mp3, .wav, .ogg, .flac, .aac, .m4a) and upload it to the server.")
bullet("Section 2 lets the user request the file back from the server, plays it with st.audio, and provides a Save File Locally download button.")
bullet("Added connection-error handling so the user gets a clear message when the server is not running.")
bullet("Used an environment variable (SERVER_URL) so the server address can be changed without editing the code.")

heading("2.4  Research on TTS and Voice Cloning Models", 12)
bullet("Studied how Text-to-Speech (TTS) and voice cloning work, including the idea of zero-shot cloning from a short reference audio sample.")
bullet("Researched leading open-source models in 2026: Coqui XTTS-v2 (clones a voice from a ~6 second clip and supports 17 languages), F5-TTS (flow-matching, clones from 5-15 seconds of audio), Chatterbox, OpenVoice, Bark, and the lightweight Kokoro-82M.")
bullet("Compared the models on cloning quality, speed, language support, and hardware (GPU/VRAM) requirements to decide which is suitable for the project.")
bullet("Concluded that XTTS-v2 and F5-TTS are the strongest candidates for the voice cloning feature because of their high-quality zero-shot cloning from very short samples.")

# ---------------- 3. API Endpoints ----------------
heading("3. API Endpoints (Backend)")
make_table([
    ["Method", "Endpoint", "Purpose"],
    ["GET", "/", "Health check - confirms the server is running."],
    ["POST", "/receive", "Receives an uploaded audio file and saves it to the uploads/ folder."],
    ["GET", "/send", "Sends the most recently uploaded audio file back to the client."],
])

# ---------------- 4. Project Workflow ----------------
heading("4. Project Workflow")
body("The application works in two simple flows between the Streamlit UI and the FastAPI server:")
bullet("Upload flow: User selects an audio file in the UI and clicks \"Upload to Server\" -> a POST request is sent to /receive -> the server saves the file and returns a success message.")
bullet("Download flow: User clicks \"Request from Server\" -> a GET request is sent to /send -> the server returns the file -> the UI plays the audio and offers a download button.")

# ---------------- 5. Technology Stack ----------------
heading("5. Technology Stack")
make_table([
    ["Layer", "Technology"],
    ["Backend", "FastAPI, Uvicorn (ASGI server)"],
    ["Frontend / UI", "Streamlit"],
    ["File Handling", "python-multipart, shutil, os, mimetypes"],
    ["HTTP Client", "requests (frontend to backend communication)"],
    ["Language", "Python 3"],
    ["Research (TTS/Cloning)", "Coqui XTTS-v2, F5-TTS, OpenVoice, Bark, Kokoro-82M"],
    ["Dev Environment", "Windows, VS Code, virtual environment (.venv)"],
])

# ---------------- 6. Skills Learned ----------------
heading("6. Skills Learned")
bullet("Building REST API endpoints with FastAPI (GET and POST, file uploads and file responses).")
bullet("Running an ASGI server with Uvicorn and testing endpoints.")
bullet("Creating an interactive frontend with Streamlit (file uploader, buttons, audio player, download button).")
bullet("Connecting a frontend to a backend over HTTP using the requests library.")
bullet("Handling files in Python: saving uploads, serving files, and creating folders automatically.")
bullet("Writing clean error handling for connection failures and missing files.")
bullet("Researching and comparing open-source TTS and voice cloning models.")

# ---------------- 7. Challenges and Solutions ----------------
heading("7. Challenges and Solutions")
make_table([
    ["Challenge", "Solution"],
    ["Frontend could not connect to the backend server",
     "Made sure the FastAPI server was running on port 8000 and added connection-error handling in Streamlit with a clear message."],
    ["Uploaded files needed a safe place to be stored",
     "Used os.makedirs(UPLOAD_DIR, exist_ok=True) so the uploads/ folder is created automatically on startup."],
    ["File upload was failing without form-data support",
     "Installed and used python-multipart so FastAPI could accept uploaded files."],
    ["Running two servers (backend + UI) at the same time",
     "Started Uvicorn in one terminal and Streamlit in another, as documented in the README."],
    ["Choosing the right TTS / voice cloning model",
     "Researched and compared multiple open-source models (XTTS-v2, F5-TTS, etc.) on quality, speed and hardware needs."],
])

# ---------------- 8. Results and Achievements ----------------
heading("8. Results and Achievements")
bullet("A working FastAPI backend running on http://localhost:8000 that can receive and send audio files.")
bullet("A Streamlit frontend that lets users upload an audio file and download it back, with audio playback.")
bullet("Audio files successfully transferred between the UI and server and stored in the uploads/ folder.")
bullet("Clear error messages when the server is offline or no file is available.")
bullet("A documented README explaining the purpose of each file and how to run the project.")
bullet("Completed research on open-source TTS and voice cloning models to guide the next phase of the project.")

# ---------------- 9. Screenshots ----------------
heading("9. Screenshots")
images = []
for ext in ("*.png", "*.jpg", "*.jpeg"):
    images.extend(sorted(glob.glob(os.path.join(SCREENSHOTS_DIR, ext))))

if images:
    for img in images:
        try:
            doc.add_picture(img, width=Inches(5.8))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cr = cap.add_run(os.path.splitext(os.path.basename(img))[0])
            cr.italic = True
            cr.font.size = Pt(9)
        except Exception as e:
            print("Skip", img, e)
else:
    doc.add_paragraph(
        "(Place screenshot images in the 'screenshots' folder and run again to embed them here.)"
    )

# ---------------- 10. Week 2 Plan ----------------
heading("10. Week 2 Plan")
bullet("Integrate a TTS model (such as Coqui XTTS-v2) so the app can convert text to speech.")
bullet("Add a basic voice cloning feature using a short reference audio sample.")
bullet("Improve the UI to support selecting a voice and entering text.")
bullet("Test the cloning quality and tune the model settings.")
bullet("Continue learning the chosen model's API and document the results.")

# ---------------- 11. Conclusion ----------------
heading("11. Conclusion")
body(
    "Week 1 successfully established the foundation of the project. A working client-server audio "
    "transfer application was built using FastAPI and Streamlit, and solid research was completed on "
    "modern open-source TTS and voice cloning models. This sets a clear path for Week 2, where the "
    "focus will shift to integrating a TTS model and implementing the core voice cloning feature."
)

# ---------------- 12. Declaration ----------------
heading("12. Declaration")
body(
    "I hereby declare that all the work described in this report was completed by me during Week 1 "
    "of my internship at Spurvance Labs. The setup, backend and frontend development, file-transfer "
    "implementation, and the research on TTS and voice cloning models were carried out by me with "
    "guidance from the project documentation and my supervisor."
)
doc.add_paragraph()
doc.add_paragraph(f"Name: {STUDENT_NAME}")
doc.add_paragraph(f"Date: {SUBMISSION_DATE}")

doc.save(OUT)
print("Saved:", OUT)
print("Screenshots embedded:", len(images))
